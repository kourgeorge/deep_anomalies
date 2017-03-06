import random
import numpy as np
from matplotlib import pyplot as plt
import tensorflow.examples.tutorials.mnist.input_data as input_data

from docx import Document
from docx.shared import Inches


mnist = input_data.read_data_sets("MNIST", one_hot=True)
mnist2 = input_data.read_data_sets("MNIST", one_hot=False)


def draw_cost_function(cost_array, images_path):
    # plot the cost graph
    plt.plot(cost_array)
    plt.ylabel('Index')
    plt.ylabel('Cost')
    plt.savefig(images_path + "cost_graph" + '.pdf')
    plt.close()

def draw_top_least_anom(cost_array, k, images_path):

    top_anom_indices = cost_array.argsort()[-k:][::-1]
    least_anom_indices = cost_array.argsort()[0:k]

    for i in range(k):
        image_vec = mnist.test.images[top_anom_indices[i]]
        image = np.reshape(image_vec, (28, -1))
        plt.imshow(image, cmap=plt.gray())
        plt.savefig(images_path + "anom_s" + str(cost_array[top_anom_indices[i]]) + '.pdf')

    for i in range(k):
        image_vec = mnist.test.images[least_anom_indices[i]]
        image = np.reshape(image_vec, (28, -1))
        plt.imshow(image, cmap=plt.gray())
        plt.savefig(images_path + "./anom_s" + str(i) + '.pdf')


def draw_riddle(cost_array,  k_anom, k_normal, images_path):
    top_anom_indices = cost_array.argsort()[-k_anom:][::-1]
    least_anom_indices = cost_array.argsort()[0:k_normal]

    indices = np.concatenate((top_anom_indices, least_anom_indices), axis=0)
    #random.shuffle(indices)

    num_columns = 3
    k = k_anom + k_normal
    num_rows = k / num_columns
    f, ax = plt.subplots(num_rows, num_columns, sharex=True)
    for i in range(k):
        # f.suptitle("Identify the following numbers!")
        image_vec = mnist.test.images[indices[i]]
        image = np.reshape(image_vec, (28, -1))
        x = i / num_columns
        y = i % num_columns
        ax[x][y].imshow(image, cmap=plt.gray())
        ax[x][y].axis('off')
        ax[x][y].set_title(str(mnist2.test.labels[indices[i]]) + "," + str(cost_array[indices[i]]))

    plt.savefig(images_path + "riddle" + '.pdf')


def compare_anom(cost_array, images_path):

    dict = sort_anomalies_by_label(cost_array)

    top_anom_indices = np.array([dict[2][0], dict[3][0], dict[5][0], dict[7][0], dict[0][0]])

    least_anom_indices = np.array([dict[2][-1], dict[3][-1], dict[5][-1], dict[7][-1], dict[0][-1]])

    indices = np.concatenate((top_anom_indices, least_anom_indices), axis=0)
    # random.shuffle(indices)

    num_columns = 5
    k = 10
    num_rows = int(k / num_columns)
    f, ax = plt.subplots(num_rows, num_columns, sharex=True)
    for i in range(k):
        # f.suptitle("Identify the following numbers!")
        image_vec = mnist.test.images[indices[i]]
        image = np.reshape(image_vec, (28, -1))
        x = int(i / num_columns)
        y = i % num_columns
        ax[x][y].imshow(image, cmap=plt.gray())
        ax[x][y].axis('off')

    plt.savefig(images_path + "compare" + '.pdf')

def draw_top_bottom_images(images, cost_array, k, images_path):

    top_anom_indices = cost_array.argsort()[-k:][::-1]
    least_anom_indices = cost_array.argsort()[0:k]

    for ind in top_anom_indices:
        image_vec = images[ind]
        image = np.reshape(image_vec, (28, -1))
        plt.imshow(image, cmap=plt.gray())
        plt.savefig(images_path + str(ind) + "-" +str(cost_array[ind]) + '.pdf')

    for ind in least_anom_indices:
        image_vec = images[ind]
        image = np.reshape(image_vec, (28, -1))
        plt.imshow(image, cmap=plt.gray())
        plt.savefig(images_path + str(ind) + "-" +str(cost_array[ind]) + '.pdf')


def draw_diffs(cost_array, original, reconstructed, images_path):
    indices = cost_array.argsort()[::-1]

    num_lines = 16
    paper_length = num_lines*4.5
    f, ax = plt.subplots(num_lines, 3, figsize=(15,paper_length), sharex=True)
    for i in range(num_lines):

        image_vec = np.power(np.subtract(original[indices[i]],reconstructed[indices[i]]),2)
        diff_image = np.reshape(image_vec, (28, -1))

        ax[i][0].imshow(diff_image, cmap=plt.hot())
        ax[i][0].axis('off')

        ax[i][1].imshow(np.reshape(original[indices[i]], (28, -1)), cmap=plt.gray())
        ax[i][1].axis('off')

        ax[i][2].imshow(np.reshape(reconstructed[indices[i]], (28,-1)), cmap=plt.gray())
        ax[i][2].axis('off')

    plt.savefig(images_path + "before_after_anom" + '.pdf')

    indices = cost_array.argsort()
    f, ax = plt.subplots(num_lines, 3, figsize=(15,paper_length), sharex=True)
    for i in range(num_lines):

        image_vec = np.power(np.subtract(original[indices[i]],reconstructed[indices[i]]),2)
        diff_image = np.reshape(image_vec, (28, -1))

        ax[i][0].imshow(diff_image, cmap=plt.hot())
        ax[i][0].axis('off')

        ax[i][1].imshow(np.reshape(original[indices[i]], (28, -1)), cmap=plt.gray())
        ax[i][1].axis('off')

        ax[i][2].imshow(np.reshape(reconstructed[indices[i]], (28,-1)), cmap=plt.gray())
        ax[i][2].axis('off')

    plt.savefig(images_path + "before_after_norm" + '.pdf')

def sort_anomalies_by_label(cost_array):
    '''return a dictionary with 10 keys for each number, containing arrays of sorted anomalies'''
    anom_sorted_indices = cost_array.argsort()
    labels = mnist2.test.labels[anom_sorted_indices]
    dict = {0: [], 1: [], 2: [], 3: [], 4: [], 5: [], 6: [], 7: [], 8: [], 9: []}

    for (label, anom_index) in zip(labels, anom_sorted_indices):
        dict[label].append(anom_index)

    return dict

def get_image_indices_of_label(number):
    train = np.where(mnist2.train.labels ==number)[0]
    test = np.where(mnist2.test.labels ==number)[0]

    return train, test


def draw_images_to_doc(images):

    image_vec = mnist.test.images[images[1]]
    image = np.reshape(image_vec, (28, -1))
    plt.imshow(image, cmap=plt.gray())

    doc = Document('addImage.docx')
    tables = doc.tables
    p = tables[0].rows[0].cells[0].add_paragraph()
    r = p.add_run()
    r.add_picture('resized.png', width=Inches(4.0), height=Inches(.7))
    p = tables[1].rows[0].cells[0].add_paragraph()
    r = p.add_run()
    r.add_picture('teste.png', width=Inches(4.0), height=Inches(.7))
    doc.save('addImage.docx')